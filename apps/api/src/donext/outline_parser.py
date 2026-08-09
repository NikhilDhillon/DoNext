import io
import re
import zipfile
from collections import Counter
from dataclasses import dataclass, field
from datetime import date, datetime, time
from pathlib import Path
from typing import Any

import pdfplumber
from docx import Document
from pypdf import PdfReader

from donext.errors import ApiError
from donext.models import (
    AllocationMethod,
    SchemeSelectionMode,
    SelectionRule,
    WeightOrigin,
)
from donext.schemas import (
    AssessmentGroupInput,
    GradingSchemeComponentInput,
    GradingSchemeInput,
    OutlineCourseProposal,
    OutlineDocumentType,
    OutlineExtractionRead,
    OutlineItemKind,
    OutlineItemProposal,
    OutlineMeetingProposal,
)

MAX_DOCUMENT_PAGES = 150
MAX_EXTRACTED_CHARACTERS = 200_000
MAX_DOCX_EXPANDED_BYTES = 25 * 1024 * 1024
MAX_DOCX_ENTRIES = 2_000

COURSE_CODE_PATTERN = re.compile(r"\b([A-Z]{2,5})\s*[- ]?\s*(\d{3,4}[A-Z]?)\b", re.I)
INSTRUCTOR_PATTERN = re.compile(
    r"(?:instructor|professor|prof\.?|lecturer)\s*[:\-]\s*([^\n|]{2,100})",
    re.IGNORECASE,
)
DATE_PATTERN = re.compile(
    r"\b(?:"
    r"(?:Mon|Tue(?:s)?|Wed|Thu(?:rs)?|Fri|Sat|Sun)(?:day)?\s*,?\s+)?(?:"
    r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
    r"Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|"
    r"Dec(?:ember)?)\s+\d{1,2}(?:st|nd|rd|th)?(?:,?\s+\d{4})?"
    r"|\d{1,2}(?:st|nd|rd|th)?\s+(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|"
    r"Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|Aug(?:ust)?|Sep(?:t(?:ember)?)?|"
    r"Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)(?:,?\s+\d{4})?"
    r"|\d{4}-\d{1,2}-\d{1,2}"
    r"|\d{1,2}/\d{1,2}(?:/\d{2,4})?"
    r")\b",
    re.IGNORECASE,
)
TIME_RANGE_PATTERN = re.compile(
    r"\b(\d{1,2}(?::\d{2})?\s*(?:a\.?m\.?|p\.?m\.?)?)\s*"
    r"(?:-|–|—|to)\s*"
    r"(\d{1,2}(?::\d{2})?\s*(?:a\.?m\.?|p\.?m\.?)?)\b",
    re.IGNORECASE,
)
WEIGHT_PATTERN = re.compile(r"\b(\d{1,3}(?:\.\d+)?)\s*%")
LOCATION_PATTERN = re.compile(r"(?:location|room)\s*[:\-]?\s*([A-Za-z0-9 .#-]{2,60})", re.I)
SEMESTER_PREFIX_PATTERN = re.compile(r"^(?:spring|summer|fall|autumn|winter)\s+\d{4}\s*", re.I)
MONTH_YEAR_PATTERN = re.compile(
    r"\b(January|February|March|April|May|June|July|August|September|October|November|December)\s+2\s*0\s*(\d)\s*(\d)\b",
    re.I,
)

ITEM_KIND_KEYWORDS: tuple[tuple[OutlineItemKind, tuple[str, ...]], ...] = (
    ("exam", ("final exam", "midterm", "exam", "test", "rewrite")),
    ("quiz", ("quiz",)),
    ("assignment", ("assignment", "homework", "problem set", "due: ass", " ass ")),
    ("project", ("project", "capstone", "sprint")),
    ("paper", ("paper", "essay", "report", "presentation")),
    ("lab", ("lab",)),
)

DAY_NAMES = {
    "monday": 0,
    "mon": 0,
    "tuesday": 1,
    "tue": 1,
    "tues": 1,
    "wednesday": 2,
    "wed": 2,
    "thursday": 3,
    "thu": 3,
    "thur": 3,
    "thurs": 3,
    "friday": 4,
    "fri": 4,
    "saturday": 5,
    "sat": 5,
    "sunday": 6,
    "sun": 6,
}
COMPACT_DAYS = {
    "m": [0],
    "mw": [0, 2],
    "mwf": [0, 2, 4],
    "mwr": [0, 2, 3],
    "mtwr": [0, 1, 2, 3],
    "mtwrf": [0, 1, 2, 3, 4],
    "wf": [2, 4],
    "t": [1],
    "tu": [1],
    "th": [3],
    "r": [3],
    "tr": [1, 3],
    "tth": [1, 3],
    "f": [4],
    "sa": [5],
    "su": [6],
}
MONTH_NUMBERS = {
    "january": 1,
    "february": 2,
    "march": 3,
    "april": 4,
    "may": 5,
    "june": 6,
    "july": 7,
    "august": 8,
    "september": 9,
    "october": 10,
    "november": 11,
    "december": 12,
}


@dataclass
class ExtractedDocument:
    text: str
    pages: list[str] = field(default_factory=list)
    tables: list[list[list[str]]] = field(default_factory=list)
    metadata: dict[str, str] = field(default_factory=dict)


def extract_outline(file_name: str, content: bytes, semester_start: date) -> OutlineExtractionRead:
    suffix = Path(file_name).suffix.lower()
    document = _extract_document(suffix, content)
    lines = _clean_lines(document.text)
    if not lines:
        raise ApiError("EMPTY_DOCUMENT", "No readable text was found in that document.", 422)

    document_type = _classify_document(file_name, document)
    course = _extract_course(lines, document, file_name)
    items = _merge_items(
        [
            *_extract_items(lines, semester_start.year),
            *_extract_items_from_tables(document.tables, semester_start.year),
            *_extract_calendar_items(document, semester_start),
        ]
    )
    grading_evidence = _grading_evidence(lines)
    items, groups, schemes = _build_grading_proposal(items, grading_evidence)
    meetings = _merge_meetings(
        [
            *_extract_meetings(lines, course),
            *_extract_meetings_from_tables(document.tables, course),
        ]
    )
    warnings = _proposal_warnings(course, items, meetings, document_type)
    if document_type == "course_schedule":
        warnings.extend(_source_date_warnings(document, semester_start))

    return OutlineExtractionRead(
        file_name=file_name,
        source_files=[file_name],
        document_types=[document_type],
        course=course,
        items=items,
        groups=groups,
        schemes=schemes,
        grading_evidence=grading_evidence,
        meetings=meetings,
        warnings=warnings,
    )


def merge_outline_extractions(
    extractions: list[OutlineExtractionRead],
) -> list[OutlineExtractionRead]:
    groups: dict[str, list[OutlineExtractionRead]] = {}
    order: list[str] = []
    for index, extraction in enumerate(extractions):
        code = _normalized_course_code(extraction.course.code)
        key = code or f"unknown:{index}"
        if key not in groups:
            groups[key] = []
            order.append(key)
        groups[key].append(extraction)

    merged: list[OutlineExtractionRead] = []
    for key in order:
        related = groups[key]
        if len(related) == 1:
            merged.append(related[0])
            continue
        course = _merge_course_proposals([item.course for item in related])
        items = _merge_items([proposal for item in related for proposal in item.items])
        grading_evidence = list(
            dict.fromkeys(evidence for item in related for evidence in item.grading_evidence)
        )
        items, assessment_groups, schemes = _build_grading_proposal(items, grading_evidence)
        meetings = _merge_meetings([proposal for item in related for proposal in item.meetings])
        source_files = [
            name for item in related for name in (item.source_files or [item.file_name])
        ]
        document_types = list(
            dict.fromkeys(kind for item in related for kind in item.document_types)
        )
        source_warnings = [
            warning
            for item in related
            for warning in item.warnings
            if warning.startswith("The document contains calendar headings")
        ]
        warnings = list(
            dict.fromkeys(
                [*_proposal_warnings(course, items, meetings, document_types[0]), *source_warnings]
            )
        )
        if len(source_files) > 1:
            warnings.insert(0, f"Combined {len(source_files)} related files for this course.")
        if "lecture_material" in document_types:
            warnings.append(
                "Lecture material supplied course and grading context; slide topics were not "
                "treated as deadlines."
            )
        merged.append(
            OutlineExtractionRead(
                file_name=source_files[0],
                source_files=source_files,
                document_types=document_types,
                course=course,
                items=items,
                groups=assessment_groups,
                schemes=schemes,
                grading_evidence=grading_evidence,
                meetings=meetings,
                warnings=warnings,
            )
        )
    return merged


def _extract_document(suffix: str, content: bytes) -> ExtractedDocument:
    if suffix == ".pdf":
        return _extract_pdf(content)
    if suffix == ".docx":
        _validate_docx_archive(content)
        try:
            document = Document(io.BytesIO(content))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            tables = [
                [[cell.text for cell in row.cells] for row in table.rows]
                for table in document.tables
            ]
            table_rows = [" | ".join(row) for table in tables for row in table]
            text = "\n".join([*paragraphs, *table_rows])
            return _checked_document(ExtractedDocument(text=text, pages=[text], tables=tables))
        except Exception as error:
            raise ApiError(
                "INVALID_DOCUMENT", "That Word document could not be read.", 422
            ) from error
    if suffix == ".txt":
        try:
            text = content.decode("utf-8-sig")
        except UnicodeDecodeError as error:
            raise ApiError(
                "INVALID_DOCUMENT", "Text documents must use UTF-8 encoding.", 422
            ) from error
        return _checked_document(ExtractedDocument(text=text, pages=[text]))
    raise ApiError("UNSUPPORTED_DOCUMENT", "Upload a PDF, DOCX, or TXT document.", 415)


def _extract_pdf(content: bytes) -> ExtractedDocument:
    try:
        reader = PdfReader(io.BytesIO(content), strict=False)
        if reader.is_encrypted:
            raise ApiError("ENCRYPTED_DOCUMENT", "Password-protected PDFs are not supported.", 422)
        if len(reader.pages) > MAX_DOCUMENT_PAGES:
            raise ApiError("DOCUMENT_TOO_LONG", "Documents may contain up to 150 pages.", 413)
        metadata = {
            str(key).removeprefix("/"): str(value)
            for key, value in (reader.metadata or {}).items()
            if value
        }
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            pages = [
                page.extract_text(layout=True) or page.extract_text() or "" for page in pdf.pages
            ]
            tables = [
                _clean_table(table)
                for page in pdf.pages
                for table in page.extract_tables()
                if table
            ]
        return _checked_document(
            ExtractedDocument(
                text="\n\f\n".join(pages), pages=pages, tables=tables, metadata=metadata
            )
        )
    except ApiError:
        raise
    except Exception as error:
        raise ApiError("INVALID_DOCUMENT", "That PDF could not be read.", 422) from error


def _checked_document(document: ExtractedDocument) -> ExtractedDocument:
    if len(document.text) > MAX_EXTRACTED_CHARACTERS:
        raise ApiError(
            "DOCUMENT_TOO_LONG", "The extracted document text is too long to process.", 413
        )
    return document


def _clean_table(table: list[list[Any]]) -> list[list[str]]:
    return [[re.sub(r"\s+", " ", str(cell or "")).strip() for cell in row] for row in table]


def _clean_lines(text: str) -> list[str]:
    return [line for raw in text.splitlines() if (line := re.sub(r"\s+", " ", raw).strip())]


def _classify_document(file_name: str, document: ExtractedDocument) -> OutlineDocumentType:
    lowered = f"{file_name}\n{document.text[:80_000]}".lower()
    headers = [{_key(cell) for cell in table[0]} for table in document.tables if table]
    if any(
        {"sunday", "monday", "tuesday", "wednesday", "thursday", "friday", "saturday"}.issubset(
            header
        )
        for header in headers
    ):
        return "course_schedule"
    if any(
        marker in lowered
        for marker in (
            "assessment methods",
            "course overview",
            "course learning outcomes",
            "course-outlines.",
        )
    ):
        return "course_outline"
    if "lecture" in file_name.lower() or re.search(r"\blecture\s+\d+\b", lowered):
        return "lecture_material"
    return "unknown"


def _extract_course(
    lines: list[str], document: ExtractedDocument, file_name: str
) -> OutlineCourseProposal:
    evidence = [file_name, *document.metadata.values(), *lines[:160]]
    weighted_codes: Counter[str] = Counter()
    for index, value in enumerate(evidence):
        for match in COURSE_CODE_PATTERN.finditer(value):
            found_code = f"{match.group(1).upper()} {match.group(2).upper()}"
            weighted_codes[found_code] += 4 if index < 1 + len(document.metadata) else 1
    code: str | None = weighted_codes.most_common(1)[0][0] if weighted_codes else None

    name_candidates: list[tuple[int, str]] = []
    if code:
        compact_code = code.replace(" ", r"\s*")
        code_pattern = re.compile(compact_code, re.I)
        for index, line in enumerate(lines[:120]):
            line_match = code_pattern.search(line)
            if not line_match:
                continue
            after = line[line_match.end() :]
            before = SEMESTER_PREFIX_PATTERN.sub("", line[: line_match.start()])
            candidate_value = after if after.strip(" :-–—") else before
            candidate_name = _clean_course_name(candidate_value)
            if candidate_name:
                score = 100 - index
                if ":" in after:
                    score += 20
                if "lecture" in candidate_name.lower():
                    score -= 25
                name_candidates.append((score, candidate_name))
        metadata_title = document.metadata.get("Title")
        if metadata_title:
            metadata_match = code_pattern.search(metadata_title)
            if metadata_match:
                metadata_name = _clean_course_name(metadata_title[metadata_match.end() :])
                if metadata_name:
                    name_candidates.append((110, metadata_name))

    name = max(name_candidates, key=lambda item: item[0])[1] if name_candidates else None
    instructor = _extract_instructor(lines, document)
    confidence = 0.95 if code and name else 0.78 if code or name else 0.3
    return OutlineCourseProposal(code=code, name=name, instructor=instructor, confidence=confidence)


def _clean_course_name(value: str) -> str | None:
    value = SEMESTER_PREFIX_PATTERN.sub("", value)
    value = re.sub(r"\(\s*Units?\s*:.*$", "", value, flags=re.I)
    value = re.sub(r"\bLecture\s+\d+.*$", "", value, flags=re.I)
    value = re.sub(r"^[\s:|\-–—]+|[\s:|\-–—]+$", "", value)
    if (
        not 3 <= len(value) <= 120
        or _looks_like_metadata(value)
        or _looks_like_month_heading(value)
    ):
        return None
    return value


def _extract_instructor(lines: list[str], document: ExtractedDocument) -> str | None:
    text = "\n".join(lines[:180])
    explicit = INSTRUCTOR_PATTERN.search(text)
    if explicit:
        return _clean_person_name(explicit.group(1))

    for index, line in enumerate(lines[:180]):
        if "instructor" not in line.lower():
            continue
        for candidate in lines[index + 1 : index + 12]:
            name_match = re.match(r"Name\s*:\s*(.{2,100})", candidate, re.I)
            if name_match:
                return _clean_person_name(name_match.group(1))

    for table in document.tables:
        if not table:
            continue
        header = [_key(cell) for cell in table[0]]
        if "instructor" in header:
            column = header.index("instructor")
            for row in table[1:]:
                if column < len(row) and row[column].strip():
                    return _clean_person_name(row[column])

    author = document.metadata.get("Author")
    if author and not any(
        value in author.lower() for value in ("department", "dept", "microsoft", "acrobat")
    ):
        return _clean_person_name(author)
    return None


def _clean_person_name(value: str) -> str:
    return re.split(
        r"\s{2,}|\s+(?:Office|Phone|Email)\s*:", value.strip(" .,-"), maxsplit=1, flags=re.I
    )[0][:120]


def _extract_items(lines: list[str], default_year: int) -> list[OutlineItemProposal]:
    proposals: list[OutlineItemProposal] = []
    for line in lines:
        lowered = f" {line.lower()} "
        if any(
            marker in lowered
            for marker in ("required to pass", "in order to pass", "score at least", "best midterm")
        ):
            continue
        summary_items = _extract_assessment_summary(line) if not DATE_PATTERN.search(line) else []
        if summary_items:
            proposals.extend(summary_items)
            continue
        kind = _item_kind(lowered)
        if kind is None or any(
            marker in lowered for marker in ("exam period", "exams begin", "exams end")
        ):
            continue
        dates = list(DATE_PATTERN.finditer(line))
        date_match = dates[-1] if dates else None
        deadline = _parse_date(date_match.group(0), default_year) if date_match else None
        if (
            deadline is None
            and not WEIGHT_PATTERN.search(line)
            and not any(marker in lowered for marker in ("tbd", "to be announced"))
        ):
            continue
        name = _normalize_item_name(_item_name(line, date_match), kind)
        weight_match = WEIGHT_PATTERN.search(line)
        weight = float(weight_match.group(1)) if weight_match else None
        if kind == "project" and name.lower().startswith("sprint"):
            weight = None
        if deadline is None and weight == 0:
            continue
        proposals.append(
            OutlineItemProposal(
                name=name,
                kind=kind,
                deadline_at=deadline,
                weight_percent=weight if weight is None or weight <= 100 else None,
                estimated_minutes=_estimated_minutes(kind),
                confidence=0.9 if deadline else 0.68,
                source_text=line[:500],
            )
        )
    return proposals[:140]


def _extract_assessment_summary(line: str) -> list[OutlineItemProposal]:
    lowered = line.lower()
    weights = [float(match.group(1)) for match in WEIGHT_PATTERN.finditer(line)]
    summaries: list[tuple[str, OutlineItemKind, float]] = []
    if "assignment" in lowered and weights:
        summaries.append(("Assignments", "assignment", weights[0]))
    if "midterm" in lowered and "final exam" in lowered and len(weights) >= 2:
        summaries.extend([("Midterms", "exam", weights[0]), ("Final Exam", "exam", weights[1])])
    if not summaries:
        return []
    return [
        OutlineItemProposal(
            name=name,
            kind=kind,
            deadline_at=None,
            weight_percent=weight,
            estimated_minutes=_estimated_minutes(kind),
            confidence=0.82,
            source_text=line[:500],
        )
        for name, kind, weight in summaries
    ]


def _extract_items_from_tables(
    tables: list[list[list[str]]], default_year: int
) -> list[OutlineItemProposal]:
    proposals: list[OutlineItemProposal] = []
    for table in tables:
        if len(table) < 2:
            continue
        header = [_key(cell) for cell in table[0]]
        name_column = _column(header, "assessment", "exams", "sprint no", "item")
        date_column = _column(header, "due date", "deadline", "date")
        weight_column = _column(header, "weight")
        if name_column is None or (date_column is None and weight_column is None):
            continue
        for row in table[1:]:
            if name_column >= len(row) or not row[name_column].strip():
                continue
            raw_name = row[name_column].strip()
            kind = _item_kind(f" {raw_name.lower()} ") or (
                "project" if "sprint" in " ".join(header) else "other"
            )
            date_value = (
                row[date_column] if date_column is not None and date_column < len(row) else ""
            )
            matches = list(DATE_PATTERN.finditer(date_value))
            deadline = _parse_date(matches[-1].group(0), default_year) if matches else None
            weight_value = (
                row[weight_column] if weight_column is not None and weight_column < len(row) else ""
            )
            weight_match = WEIGHT_PATTERN.search(weight_value)
            weight = float(weight_match.group(1)) if weight_match else None
            if "sprint no" in header:
                weight = None
            if deadline is None and (weight is None or weight == 0):
                continue
            proposals.append(
                OutlineItemProposal(
                    name=_normalize_item_name(raw_name, kind),
                    kind=kind,
                    deadline_at=deadline,
                    weight_percent=weight,
                    estimated_minutes=_estimated_minutes(kind),
                    confidence=0.96 if deadline else 0.82,
                    source_text=" | ".join(row)[:500],
                )
            )
    return proposals


def _extract_calendar_items(
    document: ExtractedDocument, semester_start: date
) -> list[OutlineItemProposal]:
    proposals: list[OutlineItemProposal] = []
    calendar_tables = [
        table
        for table in document.tables
        if table
        and [_key(cell) for cell in table[0]]
        == ["sunday", "monday", "tuesday", "wednesday", "thursday", "friday", "saturday"]
    ]
    for index, table in enumerate(calendar_tables):
        page_text = document.pages[index] if index < len(document.pages) else ""
        month_match = MONTH_YEAR_PATTERN.search(page_text)
        if not month_match:
            continue
        month = MONTH_NUMBERS[month_match.group(1).lower()]
        year = semester_start.year
        for row in table[1:]:
            for cell in row:
                cell_match = re.match(r"\s*(\d{1,2})\s*(.*)", cell, re.S)
                if not cell_match:
                    continue
                day = int(cell_match.group(1))
                details = re.sub(r"\s+", " ", cell_match.group(2)).strip()
                if not details or any(
                    marker in details.lower() for marker in ("exams begin", "exams end")
                ):
                    continue
                kind = _item_kind(f" {details.lower()} ")
                if kind is None:
                    continue
                try:
                    deadline = datetime(year, month, day, 23, 59)
                except ValueError:
                    continue
                proposals.append(
                    OutlineItemProposal(
                        name=_normalize_item_name(details, kind),
                        kind=kind,
                        deadline_at=deadline,
                        weight_percent=None,
                        estimated_minutes=_estimated_minutes(kind),
                        confidence=0.94,
                        source_text=f"{month_match.group(1)} {day}: {details}"[:500],
                    )
                )
    return proposals


def _extract_meetings(
    lines: list[str], course: OutlineCourseProposal
) -> list[OutlineMeetingProposal]:
    meetings: list[OutlineMeetingProposal] = []
    for line in lines:
        lowered = line.lower()
        if (
            "office hour" in lowered
            or "after each class" in lowered
            or lowered.startswith("time:")
            or not any(
                marker in lowered for marker in ("lecture", "class", "seminar", "tutorial", "lab")
            )
        ):
            continue
        time_match = TIME_RANGE_PATTERN.search(line)
        if time_match is None:
            continue
        if not any(marker in time_match.group(0).lower() for marker in (":", "am", "pm")):
            continue
        day_indexes = _parse_days(line[: time_match.start()]) or _parse_days(line)
        if not day_indexes:
            continue
        start_time, end_time = _parse_time_range(time_match.group(1), time_match.group(2))
        if start_time is None or end_time is None or end_time <= start_time:
            continue
        location_match = LOCATION_PATTERN.search(line)
        location = location_match.group(1).strip(" .,-") if location_match else None
        for day_index in day_indexes:
            meetings.append(_meeting(course, day_index, start_time, end_time, location, line, 0.82))
    return meetings


def _extract_meetings_from_tables(
    tables: list[list[list[str]]], course: OutlineCourseProposal
) -> list[OutlineMeetingProposal]:
    meetings: list[OutlineMeetingProposal] = []
    for table in tables:
        if len(table) < 2:
            continue
        header = [_key(cell) for cell in table[0]]
        days_column = _column(header, "days of weeks", "days", "day")
        hours_column = _column(header, "hours of day", "hours", "time")
        schedule_column = _column(header, "schedule", "meeting type", "type")
        location_column = _column(header, "location", "room")
        if days_column is None or hours_column is None:
            continue
        for row in table[1:]:
            if days_column >= len(row) or hours_column >= len(row):
                continue
            day_indexes = _parse_days(row[days_column])
            time_match = TIME_RANGE_PATTERN.search(row[hours_column])
            if not day_indexes or not time_match:
                continue
            start_time, end_time = _parse_time_range(time_match.group(1), time_match.group(2))
            if start_time is None or end_time is None or end_time <= start_time:
                continue
            location = (
                row[location_column].strip()
                if location_column is not None and location_column < len(row)
                else None
            )
            meeting_type = (
                row[schedule_column].strip()
                if schedule_column is not None and schedule_column < len(row)
                else "class"
            )
            source = " | ".join(row)
            for day_index in day_indexes:
                meetings.append(
                    _meeting(
                        course,
                        day_index,
                        start_time,
                        end_time,
                        location,
                        source,
                        0.97,
                        meeting_type,
                    )
                )
    return meetings


def _meeting(
    course: OutlineCourseProposal,
    day_index: int,
    start_time: time,
    end_time: time,
    location: str | None,
    source: str,
    confidence: float,
    meeting_type: str = "class",
) -> OutlineMeetingProposal:
    return OutlineMeetingProposal(
        title=f"{course.code or course.name or 'Course'} {meeting_type.lower()}",
        day_of_week=day_index,
        start_time=start_time,
        end_time=end_time,
        location=location or None,
        confidence=confidence,
        source_text=source[:500],
    )


def _merge_course_proposals(proposals: list[OutlineCourseProposal]) -> OutlineCourseProposal:
    code = max(
        (item for item in proposals if item.code), key=lambda item: item.confidence, default=None
    )
    name = max(
        (item for item in proposals if item.name),
        key=lambda item: (item.confidence, len(item.name or "")),
        default=None,
    )
    instructor = max(
        (item for item in proposals if item.instructor),
        key=lambda item: item.confidence,
        default=None,
    )
    return OutlineCourseProposal(
        code=code.code if code else None,
        name=name.name if name else None,
        instructor=instructor.instructor if instructor else None,
        confidence=max((item.confidence for item in proposals), default=0.3),
    )


def _merge_items(items: list[OutlineItemProposal]) -> list[OutlineItemProposal]:
    merged: list[OutlineItemProposal] = []
    for item in items:
        key = _normalized_item_name(item.name)
        existing_index = next(
            (
                index
                for index, current in enumerate(merged)
                if _normalized_item_name(current.name) == key
            ),
            None,
        )
        if existing_index is None:
            merged.append(item)
            continue
        current = merged[existing_index]
        dates = [value for value in (current.deadline_at, item.deadline_at) if value]
        close_dates = len(dates) == 2 and abs((dates[0] - dates[1]).days) <= 3
        deadline = max(dates) if close_dates else current.deadline_at or item.deadline_at
        if len(dates) == 2 and not close_dates:
            merged.append(item)
            continue
        source_text = (
            f"{current.source_text} – {item.source_text}"
            if close_dates and current.source_text != item.source_text
            else (
                current.source_text if current.confidence >= item.confidence else item.source_text
            )
        )
        merged[existing_index] = current.model_copy(
            update={
                "deadline_at": deadline,
                "weight_percent": current.weight_percent or item.weight_percent,
                "confidence": max(current.confidence, item.confidence),
                "source_text": source_text,
            }
        )
    return sorted(
        merged,
        key=lambda item: (item.deadline_at is None, item.deadline_at or datetime.max, item.name),
    )[:100]


def _grading_evidence(lines: list[str]) -> list[str]:
    markers = (
        "%",
        "drop",
        "lowest",
        "best ",
        "required to pass",
        "in order to pass",
        "score at least",
        "rewrite",
        "extra credit",
    )
    return list(
        dict.fromkeys(
            line[:1000] for line in lines if any(marker in line.lower() for marker in markers)
        )
    )[:80]


def _build_grading_proposal(
    source_items: list[OutlineItemProposal], evidence: list[str]
) -> tuple[list[OutlineItemProposal], list[AssessmentGroupInput], list[GradingSchemeInput]]:
    items = [
        item.model_copy(
            update={
                "key": None,
                "group_key": None,
                "relative_weight_percent": None,
                "weight_origin": (
                    WeightOrigin.explicit
                    if item.weight_percent is not None
                    else WeightOrigin.unknown
                ),
            }
        )
        for item in source_items
    ]
    evidence_text = "\n".join(evidence)
    groups: list[AssessmentGroupInput] = []
    group_components: list[GradingSchemeComponentInput] = []
    remove_indexes: set[int] = set()

    def add_group(
        *,
        key: str,
        name: str,
        aggregate_index: int,
        member_indexes: list[int],
        allocation: AllocationMethod = AllocationMethod.equal,
        rule: SelectionRule = SelectionRule.all,
        count: int | None = None,
    ) -> None:
        aggregate = items[aggregate_index]
        if aggregate.weight_percent is None or not member_indexes:
            return
        source = aggregate.source_text
        if rule != SelectionRule.all:
            supporting_rule = next(
                (line for line in evidence if "drop" in line.lower() or "lowest" in line.lower()),
                None,
            )
            if supporting_rule:
                source = f"{source}\n{supporting_rule}"
        groups.append(
            AssessmentGroupInput(
                key=key,
                name=name,
                allocation_method=allocation,
                weight_origin=WeightOrigin.explicit,
                extraction_confidence=aggregate.confidence,
                source_text=source,
            )
        )
        group_components.append(
            GradingSchemeComponentInput(
                target_group_key=key,
                weight_percent=aggregate.weight_percent,
                selection_rule=rule,
                selection_count=count,
            )
        )
        remove_indexes.add(aggregate_index)
        for member_index in member_indexes:
            member = items[member_index]
            relative = member.relative_weight_percent
            if allocation == AllocationMethod.explicit_percent:
                match = WEIGHT_PATTERN.search(member.source_text)
                relative = float(match.group(1)) if match else relative
            items[member_index] = member.model_copy(
                update={
                    "group_key": key,
                    "weight_percent": None,
                    "relative_weight_percent": relative,
                    "weight_origin": (
                        WeightOrigin.explicit
                        if relative is not None
                        else WeightOrigin.inferred_equal
                    ),
                }
            )

    assignment_aggregate = next(
        (
            index
            for index, item in enumerate(items)
            if item.name.lower() == "assignments"
            and item.deadline_at is None
            and item.weight_percent is not None
        ),
        None,
    )
    if assignment_aggregate is not None:
        assignment_members = [
            index
            for index, item in enumerate(items)
            if index != assignment_aggregate
            and item.kind == "assignment"
            and item.name.lower() != "assignments"
        ]
        drop_lowest = bool(
            re.search(r"lowest\s+(?:assignment\s+)?grade\s+is\s+dropped", evidence_text, re.I)
        )
        add_group(
            key="assignments",
            name="Assignments",
            aggregate_index=assignment_aggregate,
            member_indexes=assignment_members,
            rule=SelectionRule.drop_lowest_n if drop_lowest else SelectionRule.all,
            count=1 if drop_lowest else None,
        )

    midterm_aggregate = next(
        (
            index
            for index, item in enumerate(items)
            if item.name.lower() == "midterms"
            and item.deadline_at is None
            and item.weight_percent is not None
        ),
        None,
    )
    if midterm_aggregate is not None:
        midterm_members = [
            index
            for index, item in enumerate(items)
            if index != midterm_aggregate
            and "midterm" in item.name.lower()
            and item.name.lower() != "midterms"
        ]
        add_group(
            key="midterms",
            name="Midterms",
            aggregate_index=midterm_aggregate,
            member_indexes=midterm_members,
        )

    project_aggregate = next(
        (
            index
            for index, item in enumerate(items)
            if "group term project" in item.name.lower()
            and item.deadline_at is None
            and item.weight_percent is not None
        ),
        None,
    )
    if project_aggregate is not None:
        sprint_members = [
            index
            for index, item in enumerate(items)
            if index != project_aggregate
            and (item.name.lower().startswith("sprint") or "project kick-off" in item.name.lower())
        ]
        add_group(
            key="term-project",
            name="Group term project",
            aggregate_index=project_aggregate,
            member_indexes=sprint_members,
            allocation=AllocationMethod.explicit_percent,
        )

    basic_index = next(
        (index for index, item in enumerate(items) if "basic skills test" in item.name.lower()),
        None,
    )
    rewrite_index = next(
        (index for index, item in enumerate(items) if "optional rewrite" in item.name.lower()),
        None,
    )
    if basic_index is not None and rewrite_index is not None:
        basic = items[basic_index]
        minimum_match = re.search(
            r"(?:score at least|grade of)\s*(\d{1,3})%[^\n]*(?:basic skills|attempt)",
            evidence_text,
            re.I,
        )
        minimum = float(minimum_match.group(1)) if minimum_match else 60.0
        group_weight = basic.weight_percent or 0
        groups.append(
            AssessmentGroupInput(
                key="basic-skills-attempts",
                name="Basic skills test attempts",
                allocation_method=AllocationMethod.equal,
                weight_origin=WeightOrigin.explicit,
                extraction_confidence=min(basic.confidence, items[rewrite_index].confidence),
                source_text=next(
                    (line for line in evidence if "at least one attempt" in line.lower()),
                    basic.source_text,
                ),
            )
        )
        group_components.append(
            GradingSchemeComponentInput(
                target_group_key="basic-skills-attempts",
                weight_percent=group_weight,
                selection_rule=SelectionRule.highest_attempt,
                minimum_required_percent=minimum,
            )
        )
        for index in (basic_index, rewrite_index):
            item = items[index]
            items[index] = item.model_copy(
                update={
                    "group_key": "basic-skills-attempts",
                    "weight_percent": None,
                    "weight_origin": WeightOrigin.inferred_equal,
                }
            )

    items = [item for index, item in enumerate(items) if index not in remove_indexes]
    used_keys: set[str] = set()
    keyed_items: list[OutlineItemProposal] = []
    for item in items:
        base_key = re.sub(r"[^a-z0-9]+", "-", item.name.lower()).strip("-") or "item"
        key = base_key
        suffix = 2
        while key in used_keys:
            key = f"{base_key}-{suffix}"
            suffix += 1
        used_keys.add(key)
        keyed_items.append(item.model_copy(update={"key": key}))
    items = keyed_items

    direct_components = [
        GradingSchemeComponentInput(
            target_item_key=item.key,
            weight_percent=item.weight_percent,
            is_extra_credit=item.extra_credit,
            minimum_required_percent=item.minimum_required_percent,
        )
        for item in items
        if item.weight_percent is not None and item.key is not None and item.group_key is None
    ]
    standard_components = [*group_components, *direct_components]
    schemes: list[GradingSchemeInput] = []
    if standard_components:
        standard_total = sum(
            component.weight_percent
            for component in standard_components
            if not component.is_extra_credit
        )
        schemes.append(
            GradingSchemeInput(
                key="standard",
                name="Standard grading scheme",
                selection_mode=(
                    SchemeSelectionMode.best_outcome
                    if re.search(r"best midterm", evidence_text, re.I)
                    else SchemeSelectionMode.fixed
                ),
                is_primary=True,
                is_complete=abs(standard_total - 100) <= 0.01,
                components=standard_components,
            )
        )

    alternative_match = re.search(
        r"best\s+midterm\s+(\d{1,3}(?:\.\d+)?)%\s+and\s+final\s+exam\s+(\d{1,3}(?:\.\d+)?)%",
        evidence_text,
        re.I,
    )
    if alternative_match:
        alternative_components: list[GradingSchemeComponentInput] = []
        for component in standard_components:
            if component.target_group_key == "midterms":
                alternative_components.append(
                    component.model_copy(
                        update={
                            "weight_percent": float(alternative_match.group(1)),
                            "selection_rule": SelectionRule.best_n,
                            "selection_count": 1,
                        }
                    )
                )
            elif component.target_item_key and any(
                item.key == component.target_item_key and "final" in item.name.lower()
                for item in items
            ):
                alternative_components.append(
                    component.model_copy(
                        update={"weight_percent": float(alternative_match.group(2))}
                    )
                )
            else:
                alternative_components.append(component)
        alternative_total = sum(
            component.weight_percent
            for component in alternative_components
            if not component.is_extra_credit
        )
        schemes.append(
            GradingSchemeInput(
                key="best-midterm",
                name="Best-midterm alternative",
                selection_mode=SchemeSelectionMode.best_outcome,
                is_complete=abs(alternative_total - 100) <= 0.01,
                components=alternative_components,
            )
        )
    return items, groups, schemes


def _merge_meetings(meetings: list[OutlineMeetingProposal]) -> list[OutlineMeetingProposal]:
    unique: dict[tuple[int, time, time, str], OutlineMeetingProposal] = {}
    for meeting in meetings:
        key = (
            meeting.day_of_week,
            meeting.start_time,
            meeting.end_time,
            (meeting.location or "").lower(),
        )
        current = unique.get(key)
        if current is None or meeting.confidence > current.confidence:
            unique[key] = meeting
    return sorted(unique.values(), key=lambda item: (item.day_of_week, item.start_time))[:30]


def _proposal_warnings(
    course: OutlineCourseProposal,
    items: list[OutlineItemProposal],
    meetings: list[OutlineMeetingProposal],
    document_type: OutlineDocumentType,
) -> list[str]:
    warnings: list[str] = []
    if course.code is None:
        warnings.append("Course code was not found. Add it before importing.")
    if course.name is None:
        warnings.append("Course name was not found. Add it before importing.")
    if not items:
        warnings.append("No assignments or exams were found. You can add them manually.")
    else:
        undated_items = [item for item in items if item.deadline_at is None]
        if undated_items:
            count = len(undated_items)
            warnings.append(
                f"{count} academic {'item has' if count == 1 else 'items have'} no date. "
                "Add the missing date before relying on the generated schedule."
            )
    if not meetings:
        warnings.append("No recurring class times were found. Add them in the next step.")
    if document_type == "lecture_material":
        warnings.append(
            "This looks like lecture material. DoNext used its course and grading context "
            "but did not treat slide topics as deadlines."
        )
    return warnings


def _source_date_warnings(document: ExtractedDocument, semester_start: date) -> list[str]:
    source_years = {
        int(f"20{match.group(2)}{match.group(3)}")
        for match in MONTH_YEAR_PATTERN.finditer(document.text)
    }
    conflicting_years = sorted(year for year in source_years if year != semester_start.year)
    if not conflicting_years:
        return []
    listed = ", ".join(str(year) for year in conflicting_years)
    return [
        f"The document contains calendar headings for {listed}, but this semester starts in "
        f"{semester_start.year}. DoNext used {semester_start.year}; review those dates."
    ]


def _validate_docx_archive(content: bytes) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            entries = archive.infolist()
            if len(entries) > MAX_DOCX_ENTRIES:
                raise ApiError(
                    "INVALID_DOCUMENT", "That Word document contains too many files.", 422
                )
            if sum(entry.file_size for entry in entries) > MAX_DOCX_EXPANDED_BYTES:
                raise ApiError("DOCUMENT_TOO_LONG", "That Word document expands beyond 25 MB.", 413)
    except ApiError:
        raise
    except zipfile.BadZipFile as error:
        raise ApiError("INVALID_DOCUMENT", "That Word document could not be read.", 422) from error


def _parse_date(value: str, default_year: int) -> datetime | None:
    cleaned = re.sub(
        r"^(?:Mon|Tue(?:s)?|Wed|Thu(?:rs)?|Fri|Sat|Sun)(?:day)?\s*,?\s+",
        "",
        value.strip(),
        flags=re.I,
    )
    cleaned = re.sub(r"(\d)(st|nd|rd|th)\b", r"\1", cleaned, flags=re.I)
    cleaned = re.sub(r"\s+", " ", cleaned)
    has_year = bool(re.search(r"\b\d{4}\b", cleaned) or re.search(r"/\d{2}$", cleaned))
    formats = (
        (
            "%Y-%m-%d",
            "%m/%d/%Y",
            "%m/%d/%y",
            "%B %d, %Y",
            "%B %d %Y",
            "%b %d, %Y",
            "%b %d %Y",
            "%d %B %Y",
            "%d %b %Y",
        )
        if has_year
        else ("%m/%d %Y", "%B %d %Y", "%b %d %Y", "%d %B %Y", "%d %b %Y")
    )
    candidate = cleaned if has_year else f"{cleaned} {default_year}"
    for date_format in formats:
        try:
            return datetime.strptime(candidate, date_format).replace(hour=23, minute=59)
        except ValueError:
            continue
    return None


def _parse_days(value: str) -> list[int]:
    lowered = value.lower()
    matches = {
        day_index
        for name, day_index in DAY_NAMES.items()
        if re.search(rf"\b{re.escape(name)}\b", lowered)
    }
    if matches:
        return sorted(matches)
    for token in reversed(re.findall(r"\b[A-Za-z]{1,5}\b", value)):
        compact = token.lower()
        if compact in COMPACT_DAYS:
            return COMPACT_DAYS[compact]
    return []


def _parse_time_range(start_value: str, end_value: str) -> tuple[time | None, time | None]:
    start_meridiem = _meridiem(start_value)
    end_meridiem = _meridiem(end_value) or start_meridiem
    start = _parse_time(start_value, end_meridiem if start_meridiem is None else None)
    end = _parse_time(end_value, end_meridiem)
    if start and end and end <= start and _meridiem(end_value) is None and start.hour < 12:
        end = end.replace(hour=end.hour + 12 if end.hour < 12 else end.hour)
    return start, end


def _parse_time(value: str, assumed_meridiem: str | None = None) -> time | None:
    cleaned = value.lower().replace(".", "").replace(" ", "")
    meridiem = _meridiem(cleaned) or assumed_meridiem
    cleaned = cleaned.removesuffix("am").removesuffix("pm")
    try:
        parts = cleaned.split(":")
        hour = int(parts[0])
        minute = int(parts[1]) if len(parts) > 1 else 0
        if meridiem == "pm" and hour < 12:
            hour += 12
        elif meridiem == "am" and hour == 12:
            hour = 0
        return time(hour, minute)
    except (ValueError, IndexError):
        return None


def _meridiem(value: str) -> str | None:
    cleaned = value.lower().replace(".", "")
    return "pm" if "pm" in cleaned else "am" if "am" in cleaned else None


def _item_kind(value: str) -> OutlineItemKind | None:
    return next(
        (kind for kind, keywords in ITEM_KIND_KEYWORDS if any(word in value for word in keywords)),
        None,
    )


def _item_name(line: str, date_match: re.Match[str] | None) -> str:
    value = line[: date_match.start()] if date_match else line
    value = re.sub(r"\b(?:due|deadline|date)\s*[:\-]?\s*$", "", value, flags=re.I)
    value = WEIGHT_PATTERN.sub("", value)
    return re.sub(r"^[\s|:;\-–—]+|[\s|:;\-–—]+$", "", value)[:160] or "Course item"


def _normalize_item_name(value: str, kind: OutlineItemKind) -> str:
    value = re.sub(r"^\d{1,2}\s+", "", value.strip())
    value = DATE_PATTERN.sub("", value)
    value = re.sub(r"\s+(?:-|–|—|to)\s*$", "", value, flags=re.I)
    assignment = re.search(r"(?:due\s*:\s*)?ass(?:ignment)?\s*(\d+)", value, re.I)
    if assignment:
        return f"Assignment {assignment.group(1)}"
    value = re.sub(r"\s+", " ", value).strip(" .,:;-–—")
    return (value or kind.replace("_", " ").title())[:160]


def _normalized_item_name(value: str) -> str:
    value = re.sub(r"\b(?:optional|current)\b", "", value.lower())
    value = re.sub(r"[^a-z0-9]+", " ", value)
    return re.sub(r"\s+", " ", value).strip()


def _normalized_course_code(value: str | None) -> str | None:
    return re.sub(r"\s+", "", value).upper() if value else None


def _estimated_minutes(kind: str) -> int:
    return {
        "exam": 480,
        "quiz": 120,
        "assignment": 240,
        "project": 720,
        "paper": 480,
        "lab": 180,
    }.get(kind, 180)


def _key(value: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", value.lower()).strip()


def _column(header: list[str], *names: str) -> int | None:
    normalized_names = {_key(name) for name in names}
    return next((index for index, value in enumerate(header) if value in normalized_names), None)


def _looks_like_metadata(value: str) -> bool:
    lowered = value.lower()
    return any(
        marker in lowered
        for marker in ("semester", "instructor", "office", "email", "credits", "schedule", "units")
    )


def _looks_like_month_heading(value: str) -> bool:
    return bool(
        re.fullmatch(
            r"(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}",
            value,
            re.I,
        )
    )
